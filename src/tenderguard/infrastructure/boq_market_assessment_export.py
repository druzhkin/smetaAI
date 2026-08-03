# ruff: noqa: RUF001

from __future__ import annotations

from collections import Counter
from datetime import UTC
from io import BytesIO
from typing import Any
from zipfile import ZipFile

from tenderguard.application.boq_market_assessment import (
    BoqMarketTechnicalAssessmentPackage,
)
from tenderguard.domain.common import content_hash
from tenderguard.infrastructure.boq_analysis_export import _canonicalize_ooxml

_CONTENT_WIDTH_DXA = 14_400

_RELATION_LABELS = {
    "EXACT_LITERAL_NAME": "Название совпало после нормализации записи",
    "ALL_EXTRACTED_BOQ_LITERALS_PRESENT": "Все выделенные параметры ВОР присутствуют",
    "PARTIAL_LITERAL_OVERLAP": "Совпала только часть параметров",
    "NO_LITERAL_OVERLAP": "Параметры ВОР не подтверждены",
    "NO_BOQ_LITERALS": "В ВОР не выделены проверяемые параметры",
}

_BLOCKER_LABELS = {
    "MARKET_OFFER_COVERAGE_INCOMPLETE": "Не для каждой строки найдено предложение",
    "BOQ_LITERAL_GAPS_PRESENT": "В источниках отсутствует часть параметров из ВОР",
    "MARKET_SOURCE_ERRORS_PRESENT": "Есть ошибки получения рыночных страниц",
    "MARKET_STRUCTURED_DATA_FINDINGS_PRESENT": "В разметке сайтов есть неоднозначности",
    "DIAGNOSTIC_LITERAL_ASSESSMENT_NOT_GOVERNED": (
        "Буквальная проверка еще не является утвержденным сопоставлением"
    ),
    "APPROVED_TECHNICAL_ATTRIBUTE_SCHEMA_REQUIRED": (
        "Нет утвержденной схемы обязательных характеристик"
    ),
    "TECHNICAL_EQUIVALENCE_NOT_ESTABLISHED": "Техническая эквивалентность не доказана",
    "COMMERCIAL_BASIS_INCOMPLETE": "Не собраны все коммерческие условия",
    "PRICE_NORMALIZATION_REQUIRED": "Цена не приведена к единице ВОР и условиям проекта",
    "INDEPENDENT_VALIDATION_REQUIRED": "Независимый пересчет не выполнен",
    "BID_RELEASE_NOT_APPROVED": "Цена для заявки не выпущена",
}


def build_boq_market_assessment_docx(
    assessment: BoqMarketTechnicalAssessmentPackage,
) -> bytes:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError as error:  # pragma: no cover - deployment guard
        raise RuntimeError("The document-worker DOCX dependency is not installed") from error

    document = Document()
    generated_at = assessment.completed_at.astimezone(UTC).replace(tzinfo=None)
    assessment_hash = content_hash(assessment)
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.right_margin = Inches(0.5)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.5)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size, color, before, after in (
        ("Heading 1", 15, "17365D", 10, 5),
        ("Heading 2", 12, "2E74B5", 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "TenderGuard | Проверка открытых рыночных предложений"
    _set_runs(header, size=8, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Стр. ")
    _set_run(footer_run, size=8, color="666666")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    _set_run(
        title.add_run("РЫНОЧНАЯ ПРОВЕРКА ВОР"),
        size=21,
        color="17365D",
        bold=True,
    )
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    _set_run(
        subtitle.add_run(
            f"Проект {assessment.project_code} | пакет {assessment.completed_at.date().isoformat()}"
        ),
        size=13,
        color="444444",
        bold=True,
    )

    status_table = document.add_table(rows=1, cols=2)
    status_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    status_table.cell(0, 0).text = "BLOCKED"
    status_table.cell(0, 1).text = (
        "Итоговая сметная стоимость: НЕ СФОРМИРОВАНА. "
        "Найденные суммы являются сырыми ценами страниц, а не ценой единицы ВОР."
    )
    _set_table_geometry(status_table, (1800, 12_600))
    _style_table(status_table, header_rows=0, fill="FCE8E6")
    for cell in status_table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_runs(status_table.cell(0, 0).paragraphs[0], size=12, color="9B1C1C", bold=True)
    _set_runs(status_table.cell(0, 1).paragraphs[0], size=10, color="9B1C1C", bold=True)

    document.add_heading("Что система уже сделала", level=1)
    document.add_paragraph(
        "Для каждой строки ВОР система открыла заранее зафиксированные публичные страницы, "
        "сохранила исходный HTML, извлекла только структурированные предложения товара и "
        "сопоставила буквальные обозначения, размеры, классы и единицы в названиях. "
        "Результат можно воспроизвести по сохраненным доказательствам."
    )

    total_candidates = sum(len(line.candidate_assessments) for line in assessment.lines)
    exact_candidates = sum(
        item.literal_comparison.name_relation == "EXACT_LITERAL_NAME"
        for line in assessment.lines
        for item in line.candidate_assessments
    )
    full_literal_candidates = sum(
        item.literal_comparison.name_relation == "ALL_EXTRACTED_BOQ_LITERALS_PRESENT"
        for line in assessment.lines
        for item in line.candidate_assessments
    )
    lines_with_candidates = sum(bool(line.candidate_assessments) for line in assessment.lines)
    summary = document.add_table(rows=1, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fill_table(
        summary,
        ("Показатель", "Результат"),
        (
            ("Строк ВОР", str(len(assessment.lines))),
            ("Строк с найденными предложениями", str(lines_with_candidates)),
            ("Извлечено сырых предложений", str(total_candidates)),
            ("Совпадение полного названия после нормализации записи", str(exact_candidates)),
            ("Все выделенные параметры ВОР найдены", str(full_literal_candidates)),
            ("Нормализованных цен", "0"),
            ("Цена системы", "Не сформирована"),
        ),
    )
    _set_table_geometry(summary, (5200, 9200))
    _style_table(summary, header_rows=1, fill="DCE6F1")

    document.add_heading("Покрытие строк ВОР", level=1)
    coverage = document.add_table(rows=1, cols=6)
    coverage.alignment = WD_TABLE_ALIGNMENT.LEFT
    coverage_rows = tuple(
        (
            str(line.row_number),
            line.boq_description,
            line.boq_unit,
            str(line.source_observation_count),
            str(len(line.candidate_assessments)),
            _line_result(line),
        )
        for line in assessment.lines
    )
    _fill_table(
        coverage,
        ("Строка", "Наименование ВОР", "Ед.", "Страниц", "Предложений", "Результат"),
        coverage_rows,
    )
    _set_table_geometry(coverage, (700, 6500, 700, 1100, 1300, 4100))
    _style_table(coverage, header_rows=1, fill="DCE6F1")

    document.add_heading("Почему итоговая цена пока заблокирована", level=1)
    blocker_counts = Counter(assessment.global_blockers)
    blockers = document.add_table(rows=1, cols=3)
    blockers.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fill_table(
        blockers,
        ("Контроль", "Статус", "Смысл"),
        tuple(
            (code, "Не выполнен", _BLOCKER_LABELS.get(code, "Требуется проверка правила"))
            for code in blocker_counts
        ),
    )
    _set_table_geometry(blockers, (4600, 1900, 7900))
    _style_table(blockers, header_rows=1, fill="FCE8E6")

    document.add_heading("Приложение: найденные сырые предложения", level=1)
    warning = document.add_paragraph()
    _set_run(
        warning.add_run(
            "Важно: суммы ниже нельзя умножать на количество ВОР. Не подтверждены фасовка, "
            "единица товара, НДС, регион, доставка, разгрузка, оплата и техническая "
            "эквивалентность."
        ),
        size=9.5,
        color="9B1C1C",
        bold=True,
    )
    candidates = document.add_table(rows=1, cols=7)
    candidates.alignment = WD_TABLE_ALIGNMENT.LEFT
    candidate_rows: list[tuple[str, ...]] = []
    for line in assessment.lines:
        for item in line.candidate_assessments:
            source_unit = item.candidate.unit_text or item.candidate.unit_code or "не указана"
            candidate_rows.append(
                (
                    str(line.row_number),
                    line.boq_description,
                    (
                        f"{item.candidate.amount_literal} {item.candidate.currency}"
                        f"\nЕд.: {source_unit}"
                    ),
                    item.candidate.source_item_name,
                    _RELATION_LABELS[item.literal_comparison.name_relation],
                    item.source_request.display_name,
                    f"{len(item.commercial_gaps)} незаполненных условий",
                )
            )
    _fill_table(
        candidates,
        (
            "Стр.",
            "Позиция ВОР",
            "Сырая сумма",
            "Название на сайте",
            "Буквальная проверка",
            "Источник",
            "Что не хватает",
        ),
        tuple(candidate_rows),
    )
    _set_table_geometry(candidates, (700, 3000, 1300, 2800, 2100, 2300, 2200))
    _style_table(candidates, header_rows=1, fill="DCE6F1", font_size=7.5)

    document.add_heading("Реестр проверенных страниц", level=1)
    source_rows: list[tuple[str, ...]] = []
    seen_uris: set[str] = set()
    for line in assessment.lines:
        for observation in line.source_observations:
            uri = observation.request.source_uri
            if uri in seen_uris:
                continue
            seen_uris.add(uri)
            source_rows.append(
                (
                    str(line.row_number),
                    observation.request.display_name,
                    _source_observation_result(observation),
                    uri,
                )
            )
    sources = document.add_table(rows=1, cols=4)
    sources.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fill_table(
        sources,
        ("Строка ВОР", "Сайт", "Исход проверки", "Точный адрес страницы"),
        tuple(source_rows),
    )
    _set_table_geometry(sources, (1400, 3000, 2500, 7500))
    _style_table(sources, header_rows=1, fill="DCE6F1", font_size=8)

    document.add_heading("Что требуется для автоматического продолжения", level=1)
    document.add_paragraph(
        "Автоматический контур должен извлечь условия, подтвердить единицу и фасовку, "
        "применить утвержденный пересчет, добавить ФГИС ЦС и контракты, затем независимо "
        "пересчитать строки. Эксперт проверяет только готовый пакет или возвращает его на "
        "доработку."
    )

    document.add_heading("Контроль воспроизводимости", level=1)
    control = document.add_table(rows=1, cols=2)
    control.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fill_table(
        control,
        ("Идентификатор", "Значение"),
        (
            ("Хеш исходного рыночного пакета", assessment.source_market_manifest_sha256),
            ("Хеш оценки", assessment_hash),
            ("Алгоритм буквальной проверки", assessment.algorithm_version),
        ),
    )
    _set_table_geometry(control, (4600, 9800))
    _style_table(control, header_rows=1, fill="F2F4F7", font_size=8)

    document.core_properties.author = "TenderGuard"
    document.core_properties.title = f"Рыночная проверка ВОР {assessment.project_code}"
    document.core_properties.subject = "Диагностические цены и проверка наименований"
    document.core_properties.created = generated_at
    document.core_properties.modified = generated_at
    output = BytesIO()
    document.save(output)
    content = _canonicalize_ooxml(output.getvalue(), generated_at=generated_at)
    verify_boq_market_assessment_docx(content, assessment)
    return content


def verify_boq_market_assessment_docx(
    content: bytes,
    assessment: BoqMarketTechnicalAssessmentPackage,
) -> None:
    try:
        from docx import Document
    except ModuleNotFoundError as error:  # pragma: no cover - deployment guard
        raise RuntimeError("The document-worker DOCX dependency is not installed") from error

    document = Document(BytesIO(content))
    text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(cell.text for table in document.tables for row in table.rows for cell in row.cells),
        ]
    )
    required = (
        assessment.project_code,
        content_hash(assessment),
        assessment.source_market_manifest_sha256,
        "Итоговая сметная стоимость: НЕ СФОРМИРОВАНА",
        "Нормализованных цен",
    )
    if any(value not in text for value in required):
        raise ValueError("Generated market assessment DOCX lacks required control content")
    for line in assessment.lines:
        if line.boq_description not in text:
            raise ValueError("Generated market assessment DOCX omits a BoQ line")
        for observation in line.source_observations:
            if observation.request.source_uri not in text:
                raise ValueError("Generated market assessment DOCX omits a checked source page")
        for item in line.candidate_assessments:
            expected_values = (
                item.candidate.source_item_name,
                item.candidate.amount_literal,
                item.candidate.currency,
                item.source_request.source_uri,
            )
            if any(value not in text for value in expected_values):
                raise ValueError("Generated market assessment DOCX omits source evidence")
    with ZipFile(BytesIO(content), "r") as archive:
        relationship_payloads = (
            archive.read(name) for name in archive.namelist() if name.endswith(".rels")
        )
        if any(b'TargetMode="External"' in payload for payload in relationship_payloads):
            raise ValueError("Generated market assessment DOCX contains external relationships")


def _line_result(line: Any) -> str:
    if not line.candidate_assessments:
        return "BLOCKED: предложение не найдено"
    if "BOQ_LITERAL_REQUIREMENTS_NOT_PRESENT_IN_SOURCE" in line.blockers:
        return "BLOCKED: параметры ВОР подтверждены не полностью"
    return "BLOCKED: буквальная проверка пройдена, эквивалентность и цена не подтверждены"


def _source_observation_result(observation: Any) -> str:
    if observation.acquisition_error_code is not None:
        return f"Ошибка: {observation.acquisition_error_code}"
    if observation.page_result is None:
        return "BLOCKED: нет воспроизводимого результата"
    candidate_count = len(observation.page_result.candidates)
    finding_count = len(observation.page_result.extraction_findings)
    finding_suffix = f"; замечаний разметки: {finding_count}" if finding_count else ""
    if candidate_count == 0:
        return "Страница получена, структурированное предложение не найдено" + finding_suffix
    return f"Страница получена, предложений: {candidate_count}{finding_suffix}"


def _fill_table(table: Any, headers: tuple[str, ...], rows: tuple[tuple[str, ...], ...]) -> None:
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value


def _style_table(
    table: Any,
    *,
    header_rows: int,
    fill: str,
    font_size: float = 8.5,
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_properties.append(cant_split)
        if row_index < header_rows:
            table_header = OxmlElement("w:tblHeader")
            table_header.set(qn("w:val"), "true")
            row_properties.append(table_header)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cell, top=60, bottom=60, start=80, end=80)
            if row_index < header_rows or header_rows == 0:
                _shade_cell(cell, fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
                paragraph.paragraph_format.line_spacing = 1.0
                _set_runs(
                    paragraph,
                    size=font_size,
                    bold=True if row_index < header_rows else None,
                )


def _set_table_geometry(table: Any, widths_dxa: tuple[int, ...]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if len(widths_dxa) != len(table.columns) or sum(widths_dxa) != _CONTENT_WIDTH_DXA:
        raise ValueError("DOCX table widths must cover the landscape content width")
    table.autofit = False
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(_CONTENT_WIDTH_DXA))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(column_width))
        grid.append(grid_column)
    for row in table.rows:
        for cell, column_width in zip(row.cells, widths_dxa, strict=True):
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(column_width))


def _set_cell_margins(cell: Any, *, top: int, bottom: int, start: int, end: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _shade_cell(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_run(
    run: Any,
    *,
    size: float,
    color: str = "000000",
    bold: bool | None = None,
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _set_runs(
    paragraph: Any,
    *,
    size: float,
    color: str = "000000",
    bold: bool | None = None,
) -> None:
    for run in paragraph.runs:
        _set_run(run, size=size, color=color, bold=bold)
