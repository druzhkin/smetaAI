# ruff: noqa: RUF001

from __future__ import annotations

import json
import re
from collections import Counter
from datetime import UTC, datetime
from decimal import Decimal
from io import BytesIO
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from tenderguard.application.analysis_reporting import (
    BoqAnalysisReport,
    BoqAnalysisRow,
    analysis_report_hash,
)

XLSX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_REQUIRED_SHEETS = (
    "ВОР_с_оценкой",
    "Источники",
    "Сопоставления",
    "Блокировки",
    "Метаданные",
)

_BLOCKER_EXPLANATIONS = {
    "BID_RELEASE_NOT_APPROVED": "Цена проекта не прошла управляемое решение о выпуске заявки.",
    "BOQ_LINE_NOT_VERIFIED": "Строка ВОР не прошла независимую проверку.",
    "CALCULATION_SNAPSHOT_MISSING": "Нет фиксированного независимо пересчитанного снимка расчета.",
    "CONTROLLED_IMPORT_WORKFLOW_REQUIRED": "Требуется управляемый импорт исходной строки.",
    "FGIS_CS_PRICE_MISSING": "Нет проверенной и нормализованной цены ФГИС ЦС.",
    "GOVERNED_RELEASE_EXPORT_NOT_IMPLEMENTED": (
        "Текущий автономный экспорт не проверяет серверное решение о выпуске цены."
    ),
    "INDEPENDENT_ROW_REVIEW_REQUIRED": "Строка требует независимой проверки.",
    "MARKET_PRICE_MISSING": "Нет проверенной независимой рыночной цены.",
    "NOMENCLATURE_MATCH_MISSING": "Не выполнено проверяемое сопоставление номенклатуры.",
    "PRICE_DECISION_MISSING": "Нет воспроизводимого решения о предлагаемой цене.",
    "PRICE_MATRIX_NOT_AVAILABLE": "Ценовая матрица еще не сформирована.",
    "PRICE_MATRIX_ROWS_BLOCKED": "В ценовой матрице есть заблокированные строки.",
    "PRICE_POLICY_INTEGRITY_FAILED": "Нет действующей утвержденной методики цены.",
    "QUANTITY_MISSING": "Не подтверждено количество.",
    "WON_TENDER_PRICE_MISSING": "Нет доказуемой построчной цены выигранного тендера.",
}


def build_boq_analysis_workbook(report: BoqAnalysisReport) -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Side
        from openpyxl.utils import get_column_letter
    except ModuleNotFoundError as error:  # pragma: no cover - deployment guard
        raise RuntimeError("The document-worker XLSX dependency is not installed") from error

    workbook = Workbook()
    default_sheet = workbook.active
    if default_sheet is None:
        raise RuntimeError("New analysis workbook has no active worksheet")
    workbook.remove(default_sheet)
    workbook.properties.creator = "TenderGuard"
    workbook.properties.title = f"Ценовая матрица {report.project_code}"
    workbook.properties.subject = "Проверяемая ценовая матрица ВОР"
    workbook.properties.keywords = "TenderGuard, ВОР, BLOCKED, provenance"
    generated_at = report.generated_at.astimezone(UTC).replace(tzinfo=None)
    workbook.properties.created = generated_at
    workbook.properties.modified = generated_at
    workbook.calculation.fullCalcOnLoad = False
    workbook.calculation.forceFullCalc = False

    colors = {
        "navy": "17365D",
        "blue": "DCE6F1",
        "light": "F4F6F9",
        "blocked": "FCE8E6",
        "blocked_text": "9B1C1C",
        "verified": "E2F0D9",
        "verified_text": "2F6B2F",
        "muted": "666666",
        "white": "FFFFFF",
        "border": "D9E1F2",
    }
    thin = Side(style="thin", color=colors["border"])
    report_hash = analysis_report_hash(report)

    main = workbook.create_sheet(_REQUIRED_SHEETS[0])
    main_headers = (
        "№",
        "Наименование ВОР/ТЗ",
        "Ед.",
        "Количество",
        "Статус количества",
        "Код работы",
        "Каталожное сопоставление",
        "Класс сопоставления",
        "Цена выигранного тендера",
        "Наименование в тендере",
        "Источник тендера",
        "Цена ФГИС ЦС",
        "Наименование ФГИС ЦС",
        "Источник ФГИС ЦС",
        "Цена рыночная",
        "Наименование на рынке",
        "Рыночный источник",
        "Цена системы",
        "Валюта",
        "Стоимость строки",
        "Статус строки",
        "Метод выбора",
        "Комментарий и аргументация",
        "Блокировки",
        "ID строки",
        "ID позиции",
        "Дата цены",
        "Статус выпуска",
    )
    _title_block(
        main,
        title=f"Ценовая матрица проекта {report.project_code}",
        subtitle=(
            "Итоговая цена показывается только для строк с полным проверенным основанием. "
            "BLOCKED означает, что цену нельзя использовать в заявке."
        ),
        status=report.analysis_status,
        column_count=len(main_headers),
        colors=colors,
    )
    _write_header(main, row=4, headers=main_headers, colors=colors, thin=thin)
    for index, row in enumerate(report.rows, start=1):
        won = _single_source(row, "WON_TENDER")
        fgis = _single_source(row, "FGIS_CS")
        market = _single_source(row, "MARKET")
        main.append(
            [
                index,
                _safe_text(row.boq_item_name),
                _safe_text(row.boq_unit),
                row.quantity,
                _safe_text(row.quantity_status),
                _safe_text(row.work_code),
                _safe_text(row.name_match.canonical_item_id if row.name_match else None),
                _safe_text(row.name_match.match_class if row.name_match else "BLOCKED"),
                won[0],
                _safe_text(won[1]),
                _safe_text(won[2]),
                fgis[0],
                _safe_text(fgis[1]),
                _safe_text(fgis[2]),
                market[0],
                _safe_text(market[1]),
                _safe_text(market[2]),
                row.proposed_amount_per_unit,
                _safe_text(row.proposed_currency),
                row.line_amount,
                row.row_status,
                _safe_text(row.selection_method),
                _safe_text("\n".join(row.rationale)),
                _safe_text("\n".join(row.blockers)),
                _safe_text(row.row_id),
                _safe_text(row.item_id),
                row.price_as_of,
                report.release_state,
            ]
        )
    _style_data_sheet(
        main,
        header_row=4,
        data_start_row=5,
        max_row=max(4, 4 + len(report.rows)),
        max_column=len(main_headers),
        status_column=21,
        blocked_columns=(18, 20),
        colors=colors,
        thin=thin,
    )
    if report.rows:
        for column in (4, 9, 12, 15, 18, 20):
            for cell in main.iter_cols(
                min_col=column,
                max_col=column,
                min_row=5,
                max_row=main.max_row,
            ):
                for item in cell:
                    item.number_format = "#,##0.00"
        for cell in main.iter_cols(
            min_col=27,
            max_col=27,
            min_row=5,
            max_row=main.max_row,
        ):
            for item in cell:
                item.number_format = "yyyy-mm-dd"
    _set_widths(
        main,
        (
            6,
            48,
            12,
            14,
            17,
            18,
            26,
            20,
            18,
            34,
            34,
            18,
            34,
            34,
            18,
            34,
            34,
            18,
            11,
            18,
            14,
            22,
            48,
            42,
            28,
            28,
            14,
            22,
        ),
    )

    sources = workbook.create_sheet(_REQUIRED_SHEETS[1])
    source_headers = (
        "ID строки",
        "Группа",
        "Тип источника",
        "Класс доказательства",
        "Источник",
        "Наименование у источника",
        "ID записи источника",
        "URL",
        "Локатор",
        "Дата наблюдения",
        "Дата цены",
        "Действует до",
        "Наличие",
        "Срок поставки, дней",
        "Цена исходная",
        "Валюта исходная",
        "Ед. исходная",
        "Цена нормализованная",
        "Валюта нормализованная",
        "Ед. нормализованная",
        "ID нормализованной цены",
        "Хеш формулы нормализации",
        "Версия политики",
        "Технические характеристики",
        "ID наблюдения",
        "ID ревизии документа",
    )
    _sheet_heading(sources, "Паспорта источников", len(source_headers), colors)
    _write_header(sources, row=3, headers=source_headers, colors=colors, thin=thin)
    for row in report.rows:
        for source in row.sources:
            sources.append(
                [
                    _safe_text(row.row_id),
                    source.source_group,
                    source.source_type,
                    source.evidence_class,
                    _safe_text(source.display_name),
                    _safe_text(source.source_item_name),
                    _safe_text(source.source_record_id),
                    _safe_text(source.source_uri),
                    _safe_text(source.source_locator),
                    source.observed_at.isoformat(),
                    source.quote_date,
                    source.valid_until,
                    _safe_text(_availability(source.available)),
                    source.lead_time_days,
                    source.raw_amount,
                    source.raw_currency,
                    source.raw_unit,
                    source.normalized_amount_per_unit,
                    _safe_text(source.normalized_currency),
                    _safe_text(source.normalized_unit),
                    _safe_text(source.normalized_price_id),
                    _safe_text(source.normalization_formula_hash),
                    _safe_text(source.normalization_policy_version_id),
                    _safe_text(json.dumps(source.technical_attributes, ensure_ascii=False)),
                    _safe_text(source.source_observation_id),
                    _safe_text(source.source_document_revision_id),
                ]
            )
    _style_data_sheet(
        sources,
        header_row=3,
        data_start_row=4,
        max_row=max(3, sources.max_row),
        max_column=len(source_headers),
        status_column=None,
        blocked_columns=(),
        colors=colors,
        thin=thin,
    )
    if sources.max_row >= 4:
        for column in (15, 18):
            for cells in sources.iter_cols(
                min_col=column,
                max_col=column,
                min_row=4,
                max_row=sources.max_row,
            ):
                for cell in cells:
                    cell.number_format = "#,##0.00"
        for column in (10, 11, 12):
            for cells in sources.iter_cols(
                min_col=column,
                max_col=column,
                min_row=4,
                max_row=sources.max_row,
            ):
                for cell in cells:
                    cell.number_format = "yyyy-mm-dd"
    _set_widths(
        sources,
        (28, 16, 22, 22, 28, 42, 24, 44, 36, 18, 14, 14, 12, 16, 18, 12, 14, 20,
         14, 14, 28, 34, 26, 48, 30, 30),
    )

    matches = workbook.create_sheet(_REQUIRED_SHEETS[2])
    match_headers = (
        "ID строки",
        "Наименование ВОР/ТЗ",
        "Наименование/ID каталога",
        "Класс сопоставления",
        "Статус",
        "Метод",
        "Версия каталога",
        "Атрибуты ВОР/ТЗ",
        "Атрибуты каталога",
        "Различающиеся атрибуты",
        "Недостающие атрибуты",
        "Наименования тендеров",
        "Наименования ФГИС ЦС",
        "Наименования рынка",
    )
    _sheet_heading(matches, "Проверяемые сопоставления наименований", len(match_headers), colors)
    _write_header(matches, row=3, headers=match_headers, colors=colors, thin=thin)
    for row in report.rows:
        match = row.name_match
        matches.append(
            [
                _safe_text(row.row_id),
                _safe_text(row.boq_item_name),
                _safe_text(match.canonical_item_id if match else None),
                _safe_text(match.match_class if match else "BLOCKED"),
                _safe_text(match.status if match else "MISSING"),
                _safe_text(match.assessment_method if match else None),
                _safe_text(match.catalog_version_id if match else None),
                _safe_text(
                    json.dumps(match.source_attributes, ensure_ascii=False) if match else None
                ),
                _safe_text(
                    json.dumps(match.canonical_attributes, ensure_ascii=False) if match else None
                ),
                _safe_text("\n".join(match.mismatched_attributes) if match else None),
                _safe_text("\n".join(match.missing_attributes) if match else None),
                _safe_text(_source_names(row, "WON_TENDER")),
                _safe_text(_source_names(row, "FGIS_CS")),
                _safe_text(_source_names(row, "MARKET")),
            ]
        )
    _style_data_sheet(
        matches,
        header_row=3,
        data_start_row=4,
        max_row=max(3, matches.max_row),
        max_column=len(match_headers),
        status_column=5,
        blocked_columns=(),
        colors=colors,
        thin=thin,
    )
    _set_widths(matches, (28, 48, 34, 22, 16, 22, 24, 44, 44, 30, 30, 40, 40, 40))

    blockers = workbook.create_sheet(_REQUIRED_SHEETS[3])
    blocker_headers = ("ID строки", "Наименование ВОР/ТЗ", "Код блокировки", "Пояснение")
    _sheet_heading(blockers, "Блокировки и отсутствующие доказательства", 4, colors)
    _write_header(blockers, row=3, headers=blocker_headers, colors=colors, thin=thin)
    for code in report.global_blockers:
        blockers.append(["PROJECT", report.project_code, code, _blocker_explanation(code)])
    for row in report.rows:
        for code in row.blockers:
            blockers.append(
                [
                    _safe_text(row.row_id),
                    _safe_text(row.boq_item_name),
                    _safe_text(code),
                    _safe_text(_blocker_explanation(code)),
                ]
            )
    _style_data_sheet(
        blockers,
        header_row=3,
        data_start_row=4,
        max_row=max(3, blockers.max_row),
        max_column=4,
        status_column=None,
        blocked_columns=(),
        colors=colors,
        thin=thin,
    )
    _set_widths(blockers, (28, 48, 42, 62))

    metadata = workbook.create_sheet(_REQUIRED_SHEETS[4])
    _sheet_heading(metadata, "Метаданные и контроль целостности", 2, colors)
    _write_header(metadata, row=3, headers=("Поле", "Значение"), colors=colors, thin=thin)
    metadata_rows: tuple[tuple[str, Any], ...] = (
        ("schema_version", report.schema_version),
        ("report_content_hash", report_hash),
        ("project_id", report.project_id),
        ("project_code", report.project_code),
        ("source_kind", report.source_kind),
        ("source_reference", report.source_reference),
        ("source_content_hash", report.source_content_hash),
        ("generated_at", report.generated_at.isoformat()),
        ("analysis_status", report.analysis_status),
        ("release_state", report.release_state),
        ("row_count", len(report.rows)),
        ("blocked_row_count", report.blocked_row_count),
        ("calculation_snapshot_id", report.calculation_snapshot_id),
        ("final_total", report.final_total),
        ("final_currency", report.final_currency),
        ("global_blockers", "\n".join(report.global_blockers)),
        ("warnings", "\n".join(report.warnings)),
    )
    for key, value in metadata_rows:
        metadata.append([key, _safe_text(value) if isinstance(value, str) else value])
    _style_data_sheet(
        metadata,
        header_row=3,
        data_start_row=4,
        max_row=metadata.max_row,
        max_column=2,
        status_column=None,
        blocked_columns=(),
        colors=colors,
        thin=thin,
    )
    _set_widths(metadata, (34, 110))

    for sheet in workbook.worksheets:
        from openpyxl.worksheet.properties import PageSetupProperties

        sheet.sheet_view.showGridLines = False
        sheet.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)
        sheet.page_setup.fitToWidth = 1
        sheet.page_setup.fitToHeight = 0
        sheet.page_margins.left = 0.25
        sheet.page_margins.right = 0.25
        sheet.page_margins.top = 0.5
        sheet.page_margins.bottom = 0.5
        sheet.freeze_panes = "A5" if sheet is main else "A4"
        if sheet.max_row >= (4 if sheet is main else 3):
            header_row = 4 if sheet is main else 3
            sheet.auto_filter.ref = (
                f"A{header_row}:{get_column_letter(sheet.max_column)}{sheet.max_row}"
            )

    output = BytesIO()
    workbook.save(output)
    workbook.close()
    content = _canonicalize_ooxml(output.getvalue(), generated_at=generated_at)
    verify_boq_analysis_workbook(content, report)
    return content


def verify_boq_analysis_workbook(content: bytes, report: BoqAnalysisReport) -> None:
    try:
        from openpyxl import load_workbook
    except ModuleNotFoundError as error:  # pragma: no cover - deployment guard
        raise RuntimeError("The document-worker XLSX dependency is not installed") from error

    workbook = load_workbook(BytesIO(content), data_only=False, read_only=False, keep_links=True)
    try:
        if tuple(workbook.sheetnames) != _REQUIRED_SHEETS:
            raise ValueError("Generated analysis workbook has an unexpected sheet set")
        if any(sheet.sheet_state != "visible" for sheet in workbook.worksheets):
            raise ValueError("Generated analysis workbook contains hidden sheets")
        if getattr(workbook, "_external_links", []):
            raise ValueError("Generated analysis workbook contains external links")
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                if any(cell.data_type == "f" for cell in row):
                    raise ValueError("Generated analysis workbook contains formulas")
        main = workbook[_REQUIRED_SHEETS[0]]
        headers: dict[str, int] = {}
        for cell in main[4]:
            if isinstance(cell.value, str) and isinstance(cell.column, int):
                headers[cell.value] = cell.column
        required_headers = {
            "ID строки",
            "Цена системы",
            "Стоимость строки",
            "Статус строки",
        }
        if not required_headers.issubset(headers):
            raise ValueError("Generated analysis workbook main sheet is incomplete")
        if main.max_row != 4 + len(report.rows):
            raise ValueError("Generated analysis workbook row count differs from its report")
        expected = {row.row_id: row for row in report.rows}
        for row_number in range(5, main.max_row + 1):
            raw_row_id = main.cell(row_number, headers["ID строки"]).value
            row_id = raw_row_id if isinstance(raw_row_id, str) else None
            if row_id is None:
                raise ValueError("Generated analysis workbook row has no string identity")
            report_row = expected.get(row_id)
            if report_row is None:
                raise ValueError("Generated analysis workbook contains an unknown row")
            if main.cell(row_number, headers["Статус строки"]).value != report_row.row_status:
                raise ValueError("Generated analysis workbook row status differs from its report")
            if report_row.row_status == "BLOCKED" and (
                main.cell(row_number, headers["Цена системы"]).value is not None
                or main.cell(row_number, headers["Стоимость строки"]).value is not None
            ):
                raise ValueError("Generated analysis workbook leaks a blocked financial value")
        metadata = workbook[_REQUIRED_SHEETS[4]]
        metadata_map = {
            metadata.cell(row, 1).value: metadata.cell(row, 2).value
            for row in range(4, metadata.max_row + 1)
        }
        if metadata_map.get("report_content_hash") != analysis_report_hash(report):
            raise ValueError("Generated analysis workbook report hash differs")
    finally:
        workbook.close()


def build_boq_analysis_docx(report: BoqAnalysisReport) -> bytes:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError as error:  # pragma: no cover - deployment guard
        raise RuntimeError("The document-worker DOCX dependency is not installed") from error

    del WD_SECTION  # retained in the guarded import so python-docx surface is checked together
    document = Document()
    generated_at = report.generated_at.astimezone(UTC).replace(tzinfo=None)
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "TenderGuard | Аналитический отчет по ВОР"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_docx_runs(header, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Стр. ")
    _set_docx_run(footer_run, size=9, color="666666")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("ОТЧЕТ ПО ЦЕНОВОЙ МАТРИЦЕ")
    _set_docx_run(title_run, size=23, color="17365D", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(f"Проект {report.project_code}")
    _set_docx_run(subtitle_run, size=14, color="444444", bold=True)

    for label, value in (
        ("Дата формирования", report.generated_at.date().isoformat()),
        ("Источник", report.source_reference),
        ("Статус выпуска", report.release_state),
        ("Контрольный хеш", analysis_report_hash(report)),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        _set_docx_run(label_run, size=10, bold=True)
        value_run = paragraph.add_run(value)
        _set_docx_run(value_run, size=10)

    document.add_paragraph().paragraph_format.space_after = Pt(2)
    status_table = document.add_table(rows=1, cols=1)
    status_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    status_cell = status_table.cell(0, 0)
    status_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    status_text = (
        "BLOCKED — итоговая сметная цена не выпущена"
        if report.analysis_status == "BLOCKED"
        else "APPROVED_FOR_BID — выпущено по фиксированному расчету"
    )
    status_paragraph = status_cell.paragraphs[0]
    status_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    status_run = status_paragraph.add_run(status_text)
    _set_docx_run(
        status_run,
        size=12,
        bold=True,
        color="9B1C1C" if report.analysis_status == "BLOCKED" else "2F6B2F",
    )
    _set_table_geometry(status_table, (9360,), qn=qn, OxmlElement=OxmlElement)
    _shade_cell(status_cell, "FCE8E6" if report.analysis_status == "BLOCKED" else "E2F0D9")

    lead = document.add_paragraph()
    lead.paragraph_format.space_before = Pt(8)
    lead.paragraph_format.space_after = Pt(8)
    lead.add_run(
        "Система сформировала проверяемый пакет, но не подставляет цену там, где "
        "нет подтвержденного сопоставления, обязательного источника или выпуска расчета."
    )

    document.add_heading("Результат обработки", level=1)
    summary_rows = (
        ("Строк ВОР", str(len(report.rows))),
        ("Строк с ценой системы", str(sum(row.row_status == "VERIFIED" for row in report.rows))),
        ("Заблокированных строк", str(report.blocked_row_count)),
        (
            "Итоговая стоимость",
            (
                f"{_decimal_text(report.final_total)} {report.final_currency}"
                if report.final_total is not None
                else "Не сформирована"
            ),
        ),
    )
    summary = document.add_table(rows=1, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fill_table(summary, ("Показатель", "Результат"), summary_rows)
    _set_table_geometry(summary, (3000, 6360), qn=qn, OxmlElement=OxmlElement)
    _style_docx_table(summary, header_fill="DCE6F1")

    document.add_heading("Покрытие источниками", level=1)
    coverage_rows = []
    for group, label in (
        ("WON_TENDER", "Выигранные тендеры"),
        ("FGIS_CS", "ФГИС ЦС"),
        ("MARKET", "Рынок"),
    ):
        rows_with_source = sum(
            any(source.source_group == group for source in row.sources)
            for row in report.rows
        )
        normalized = sum(
            source.normalized_amount_per_unit is not None
            for row in report.rows
            for source in row.sources
            if source.source_group == group
        )
        coverage_rows.append((label, str(rows_with_source), str(normalized)))
    coverage = document.add_table(rows=1, cols=3)
    coverage.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fill_table(
        coverage,
        ("Источник", "Строк с источником", "Нормализованных цен"),
        tuple(coverage_rows),
    )
    _set_table_geometry(coverage, (3600, 2880, 2880), qn=qn, OxmlElement=OxmlElement)
    _style_docx_table(coverage, header_fill="DCE6F1")

    document.add_heading("Основные причины блокировки", level=1)
    counts = Counter(
        (*report.global_blockers, *(code for row in report.rows for code in row.blockers))
    )
    blocker_rows = tuple(
        (code, str(count), _blocker_explanation(code))
        for code, count in counts.most_common(10)
    ) or (("Нет", "0", "Обязательные блокировки отсутствуют."),)
    blocker_table = document.add_table(rows=1, cols=3)
    blocker_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fill_table(blocker_table, ("Код", "Кол-во", "Что означает"), blocker_rows)
    _set_table_geometry(blocker_table, (2700, 1100, 5560), qn=qn, OxmlElement=OxmlElement)
    _style_docx_table(blocker_table, header_fill="FCE8E6")

    document.add_heading("Как использовать результат", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run(
        "В Excel показаны исходное наименование ВОР, названия у каждого источника, "
        "цены, ссылки, сопоставление и причины блокировки. Строки BLOCKED возвращаются "
        "в автоматическую обработку. Эксперт в конце проверяет готовый пакет и не вводит "
        "новую цену вручную. Если эксперт изменяет существенное сопоставление или цену, "
        "это решение требует независимого подтверждения."
    )

    document.add_heading("Следующее действие", level=1)
    next_action = document.add_paragraph()
    if report.analysis_status == "BLOCKED":
        next_action.add_run(
            "Получить отсутствующие проверенные источники, завершить сопоставление, "
            "выполнить независимый пересчет и заново сформировать пакет. До этого итог "
            "нельзя использовать как цену заявки."
        )
    else:
        next_action.add_run(
            "Эксперт сверяет контрольный хеш и выпускает только зафиксированный результат."
        )

    document.core_properties.author = "TenderGuard"
    document.core_properties.title = f"Отчет по ценовой матрице {report.project_code}"
    document.core_properties.subject = "Контролируемый итог расчета ВОР"
    document.core_properties.created = generated_at
    document.core_properties.modified = generated_at
    output = BytesIO()
    document.save(output)
    content = _canonicalize_ooxml(output.getvalue(), generated_at=generated_at)
    verify_boq_analysis_docx(content, report)
    return content


def verify_boq_analysis_docx(content: bytes, report: BoqAnalysisReport) -> None:
    try:
        from docx import Document
    except ModuleNotFoundError as error:  # pragma: no cover - deployment guard
        raise RuntimeError("The document-worker DOCX dependency is not installed") from error
    document = Document(BytesIO(content))
    text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )
    required = (
        report.project_code,
        analysis_report_hash(report),
        report.analysis_status,
        "Итоговая стоимость",
    )
    if any(value not in text for value in required):
        raise ValueError("Generated analysis DOCX is missing required control content")
    if report.analysis_status == "BLOCKED" and report.final_total is not None:
        raise ValueError("Blocked analysis report unexpectedly contains a final total")


def _title_block(
    sheet: Any,
    *,
    title: str,
    subtitle: str,
    status: str,
    column_count: int,
    colors: dict[str, str],
) -> None:
    from openpyxl.styles import Alignment, Font, PatternFill

    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=column_count)
    sheet.cell(1, 1, title)
    sheet.cell(1, 1).font = Font(bold=True, size=16, color=colors["white"])
    sheet.cell(1, 1).fill = PatternFill("solid", fgColor=colors["navy"])
    sheet.cell(1, 1).alignment = Alignment(vertical="center")
    sheet.row_dimensions[1].height = 28
    sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=column_count - 3)
    sheet.cell(2, 1, subtitle)
    sheet.cell(2, 1).alignment = Alignment(wrap_text=True, vertical="center")
    sheet.cell(2, 1).fill = PatternFill("solid", fgColor=colors["light"])
    sheet.merge_cells(
        start_row=2,
        start_column=column_count - 2,
        end_row=2,
        end_column=column_count,
    )
    status_cell = sheet.cell(2, column_count - 2, status)
    status_cell.font = Font(
        bold=True,
        color=colors["blocked_text"] if status == "BLOCKED" else colors["verified_text"],
    )
    status_cell.fill = PatternFill(
        "solid",
        fgColor=colors["blocked"] if status == "BLOCKED" else colors["verified"],
    )
    status_cell.alignment = Alignment(horizontal="center", vertical="center")
    sheet.row_dimensions[2].height = 42


def _sheet_heading(sheet: Any, title: str, column_count: int, colors: dict[str, str]) -> None:
    from openpyxl.styles import Alignment, Font, PatternFill

    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=column_count)
    cell = sheet.cell(1, 1, title)
    cell.font = Font(bold=True, size=15, color=colors["white"])
    cell.fill = PatternFill("solid", fgColor=colors["navy"])
    cell.alignment = Alignment(vertical="center")
    sheet.row_dimensions[1].height = 26


def _write_header(
    sheet: Any,
    *,
    row: int,
    headers: tuple[str, ...],
    colors: dict[str, str],
    thin: Any,
) -> None:
    from openpyxl.styles import Alignment, Border, Font, PatternFill

    for column, header in enumerate(headers, start=1):
        cell = sheet.cell(row, column, header)
        cell.font = Font(bold=True, color=colors["navy"])
        cell.fill = PatternFill("solid", fgColor=colors["blue"])
        cell.border = Border(bottom=thin)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    sheet.row_dimensions[row].height = 50


def _style_data_sheet(
    sheet: Any,
    *,
    header_row: int,
    data_start_row: int,
    max_row: int,
    max_column: int,
    status_column: int | None,
    blocked_columns: tuple[int, ...],
    colors: dict[str, str],
    thin: Any,
) -> None:
    from openpyxl.styles import Alignment, Border, Font, PatternFill

    if max_row >= data_start_row:
        for row in sheet.iter_rows(
            min_row=data_start_row,
            max_row=max_row,
            min_col=1,
            max_col=max_column,
        ):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(bottom=thin)
                cell.font = Font(size=10)
            if status_column is not None:
                status = row[status_column - 1].value
                if status in {"BLOCKED", "MISSING"}:
                    row[status_column - 1].fill = PatternFill("solid", fgColor=colors["blocked"])
                    row[status_column - 1].font = Font(
                        bold=True,
                        color=colors["blocked_text"],
                    )
                    for column in blocked_columns:
                        if row[column - 1].value is not None:
                            raise ValueError("Blocked workbook row contains a financial value")
                elif status == "VERIFIED":
                    row[status_column - 1].fill = PatternFill("solid", fgColor=colors["verified"])
                    row[status_column - 1].font = Font(
                        bold=True,
                        color=colors["verified_text"],
                    )
    sheet.auto_filter.ref = f"A{header_row}:{sheet.cell(header_row, max_column).coordinate}"


def _set_widths(sheet: Any, widths: tuple[float, ...]) -> None:
    from openpyxl.utils import get_column_letter

    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width


def _single_source(row: BoqAnalysisRow, group: str) -> tuple[Decimal | None, str, str]:
    sources = tuple(source for source in row.sources if source.source_group == group)
    normalized = tuple(
        source for source in sources if source.normalized_amount_per_unit is not None
    )
    names = _source_names(row, group)
    locations = "\n".join(
        dict.fromkeys(source.source_uri or source.source_locator for source in sources)
    )
    if len(normalized) != 1:
        return None, names, locations
    return normalized[0].normalized_amount_per_unit, names, locations


def _source_names(row: BoqAnalysisRow, group: str) -> str:
    return "\n".join(
        dict.fromkeys(
            source.source_item_name for source in row.sources if source.source_group == group
        )
    )


def _safe_text(value: Any) -> str | None:
    if value is None:
        return None
    rendered = str(value).replace("\x00", "")
    if len(rendered) > 32_767:
        rendered = rendered[:32_740] + "… [сокращено]"
    if rendered.startswith(("=", "+", "-", "@")):
        return "'" + rendered
    return rendered


def _availability(value: bool | None) -> str:
    if value is True:
        return "Да"
    if value is False:
        return "Нет"
    return "Не подтверждено"


def _blocker_explanation(code: str) -> str:
    base = code.split(":", 1)[0]
    return _BLOCKER_EXPLANATIONS.get(
        base,
        f"Не выполнено контрольное требование {base}.",
    )


def _decimal_text(value: Decimal | None) -> str:
    if value is None:
        return ""
    return f"{value:,.2f}".replace(",", " ")


def _canonicalize_ooxml(content: bytes, *, generated_at: datetime) -> bytes:
    source = BytesIO(content)
    target = BytesIO()
    with ZipFile(source, "r") as source_archive, ZipFile(
        target,
        "w",
        compression=ZIP_DEFLATED,
        compresslevel=9,
    ) as target_archive:
        for source_info in sorted(source_archive.infolist(), key=lambda item: item.filename):
            payload = source_archive.read(source_info.filename)
            if source_info.filename == "docProps/core.xml":
                timestamp = generated_at.strftime("%Y-%m-%dT%H:%M:%SZ").encode()
                for field in (b"created", b"modified"):
                    pattern = (
                        rb"(<dcterms:"
                        + field
                        + rb"\b[^>]*>)[^<]*(</dcterms:"
                        + field
                        + rb">)"
                    )
                    timestamp_match = re.search(pattern, payload)
                    if timestamp_match is None:
                        raise ValueError("Generated OOXML core timestamps are incomplete")
                    replacement = timestamp_match.group(1) + timestamp + timestamp_match.group(2)
                    payload = (
                        payload[: timestamp_match.start()]
                        + replacement
                        + payload[timestamp_match.end() :]
                    )
            target_info = ZipInfo(
                filename=source_info.filename,
                date_time=(1980, 1, 1, 0, 0, 0),
            )
            target_info.compress_type = ZIP_DEFLATED
            target_info.create_system = source_info.create_system
            target_info.external_attr = source_info.external_attr
            target_info.internal_attr = source_info.internal_attr
            target_info.flag_bits = source_info.flag_bits
            target_archive.writestr(
                target_info,
                payload,
                compress_type=ZIP_DEFLATED,
                compresslevel=9,
            )
    return target.getvalue()


def _set_docx_run(
    run: Any,
    *,
    size: float,
    color: str = "000000",
    bold: bool | None = None,
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _set_docx_runs(paragraph: Any, *, size: float, color: str) -> None:
    for run in paragraph.runs:
        _set_docx_run(run, size=size, color=color)


def _fill_table(table: Any, headers: tuple[str, ...], rows: tuple[tuple[str, ...], ...]) -> None:
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value


def _style_docx_table(table: Any, *, header_fill: str) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
            if row_index == 0:
                _shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    _set_docx_run(run, size=9.5, bold=True if row_index == 0 else None)


def _set_table_geometry(
    table: Any,
    widths_dxa: tuple[int, ...],
    *,
    qn: Any,
    OxmlElement: Any,
) -> None:
    if len(widths_dxa) != len(table.columns) or sum(widths_dxa) != 9360:
        raise ValueError("DOCX table widths must cover the exact 9360 DXA content width")
    table.autofit = False
    table.alignment = 0
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), "9360")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(column_width))
        grid.append(grid_column)
    for row in table.rows:
        for cell, column_width in zip(row.cells, widths_dxa, strict=True):
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(column_width))


def _set_cell_margins(cell: Any, *, top: int, bottom: int, start: int, end: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _shade_cell(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
